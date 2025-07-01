"""
Knowledge Management System services for DeGeNz Lounge.
This module provides services for managing knowledge repositories, items, and related functionality.
"""

import os
import logging
import tempfile
from typing import List, Dict, Any, Optional
from datetime import datetime
from sqlalchemy.orm import Session
from fastapi import UploadFile, HTTPException
import numpy as np

from app.models.knowledge_models import (
    KnowledgeRepository, 
    KnowledgeItem, 
    KnowledgeTag, 
    KnowledgeConnection,
    Document,
    Citation,
    WebSearchResult
)
from app.services.ai.unified_service import UnifiedAIService

# Initialize logging
logger = logging.getLogger(__name__)

# Initialize AI service for embeddings and processing
ai_service = UnifiedAIService()

class KnowledgeService:
    """Service for managing knowledge repositories and items."""
    
    @staticmethod
    async def create_repository(db: Session, name: str, description: str, owner_id: int, is_public: bool = False) -> KnowledgeRepository:
        """Create a new knowledge repository."""
        repository = KnowledgeRepository(
            name=name,
            description=description,
            owner_id=owner_id,
            is_public=is_public
        )
        db.add(repository)
        db.commit()
        db.refresh(repository)
        return repository
    
    @staticmethod
    async def get_repositories(db: Session, user_id: int, include_public: bool = True) -> List[KnowledgeRepository]:
        """Get all repositories accessible by a user."""
        query = db.query(KnowledgeRepository).filter(
            (KnowledgeRepository.owner_id == user_id) | 
            (KnowledgeRepository.shared_users.any(id=user_id))
        )
        
        if include_public:
            query = query.union(
                db.query(KnowledgeRepository).filter(KnowledgeRepository.is_public == True)
            )
            
        return query.all()
    
    @staticmethod
    async def get_repository(db: Session, repository_id: int, user_id: int) -> Optional[KnowledgeRepository]:
        """Get a specific repository if accessible by the user."""
        repository = db.query(KnowledgeRepository).filter(KnowledgeRepository.id == repository_id).first()
        
        if not repository:
            return None
            
        # Check if user has access
        if (repository.owner_id == user_id or 
            repository.is_public or 
            any(user.id == user_id for user in repository.shared_users)):
            return repository
            
        return None
    
    @staticmethod
    async def share_repository(db: Session, repository_id: int, owner_id: int, user_ids: List[int]) -> KnowledgeRepository:
        """Share a repository with other users."""
        repository = db.query(KnowledgeRepository).filter(
            KnowledgeRepository.id == repository_id,
            KnowledgeRepository.owner_id == owner_id
        ).first()
        
        if not repository:
            raise HTTPException(status_code=404, detail="Repository not found or you don't have permission")
        
        # Get users to share with
        from app.models.models import User
        users = db.query(User).filter(User.id.in_(user_ids)).all()
        
        # Add users to shared_users
        for user in users:
            if user not in repository.shared_users and user.id != owner_id:
                repository.shared_users.append(user)
        
        db.commit()
        db.refresh(repository)
        return repository
    
    @staticmethod
    async def create_knowledge_item(
        db: Session, 
        repository_id: int, 
        title: str, 
        content: str, 
        creator_id: int,
        content_type: str = "text",
        source_url: Optional[str] = None,
        source_type: str = "user",
        importance: float = 1.0,
        metadata: Optional[Dict[str, Any]] = None,
        tag_ids: Optional[List[int]] = None
    ) -> KnowledgeItem:
        """Create a new knowledge item."""
        # Check repository access
        repository = await KnowledgeService.get_repository(db, repository_id, creator_id)
        if not repository:
            raise HTTPException(status_code=404, detail="Repository not found or you don't have permission")
        
        # Generate embedding for the content
        embedding = await KnowledgeService._generate_embedding(content)
        
        # Create knowledge item
        knowledge_item = KnowledgeItem(
            title=title,
            content=content,
            content_type=content_type,
            source_url=source_url,
            source_type=source_type,
            importance=importance,
            metadata=metadata or {},
            embedding=embedding,
            repository_id=repository_id,
            creator_id=creator_id
        )
        
        db.add(knowledge_item)
        db.commit()
        
        # Add tags if provided
        if tag_ids:
            tags = db.query(KnowledgeTag).filter(KnowledgeTag.id.in_(tag_ids)).all()
            knowledge_item.tags = tags
            db.commit()
            
        db.refresh(knowledge_item)
        return knowledge_item
    
    @staticmethod
    async def search_knowledge(
        db: Session, 
        user_id: int, 
        query: str, 
        repository_id: Optional[int] = None,
        tag_ids: Optional[List[int]] = None,
        limit: int = 10
    ) -> List[KnowledgeItem]:
        """Search knowledge items using semantic search."""
        # Generate embedding for the query
        query_embedding = await KnowledgeService._generate_embedding(query)
        
        # Get accessible repositories
        accessible_repos = await KnowledgeService.get_repositories(db, user_id)
        accessible_repo_ids = [repo.id for repo in accessible_repos]
        
        if repository_id and repository_id not in accessible_repo_ids:
            raise HTTPException(status_code=403, detail="You don't have access to this repository")
        
        # Filter by repository if specified
        repo_filter = KnowledgeItem.repository_id.in_(accessible_repo_ids)
        if repository_id:
            repo_filter = KnowledgeItem.repository_id == repository_id
        
        # Base query
        knowledge_items = db.query(KnowledgeItem).filter(repo_filter)
        
        # Filter by tags if specified
        if tag_ids:
            for tag_id in tag_ids:
                knowledge_items = knowledge_items.filter(KnowledgeItem.tags.any(id=tag_id))
        
        # Get all items that match the filters
        items = knowledge_items.all()
        
        # Calculate similarity scores
        scored_items = []
        for item in items:
            if item.embedding:
                similarity = KnowledgeService._calculate_similarity(query_embedding, item.embedding)
                scored_items.append((item, similarity))
        
        # Sort by similarity and return top results
        scored_items.sort(key=lambda x: x[1], reverse=True)
        return [item for item, score in scored_items[:limit]]
    
    @staticmethod
    async def _generate_embedding(text: str) -> List[float]:
        """Generate embedding vector for text using AI service."""
        try:
            # Use the default AI provider for embeddings
            response = ai_service.generate_embedding(text)
            return response.get("embedding", [])
        except Exception as e:
            logger.error(f"Error generating embedding: {e}")
            # Return empty embedding if failed
            return []
    
    @staticmethod
    def _calculate_similarity(embedding1: List[float], embedding2: List[float]) -> float:
        """Calculate cosine similarity between two embeddings."""
        if not embedding1 or not embedding2:
            return 0.0
            
        try:
            # Convert to numpy arrays
            vec1 = np.array(embedding1)
            vec2 = np.array(embedding2)
            
            # Calculate cosine similarity
            dot_product = np.dot(vec1, vec2)
            norm1 = np.linalg.norm(vec1)
            norm2 = np.linalg.norm(vec2)
            
            if norm1 == 0 or norm2 == 0:
                return 0.0
                
            return dot_product / (norm1 * norm2)
        except Exception as e:
            logger.error(f"Error calculating similarity: {e}")
            return 0.0
    
    @staticmethod
    async def create_connection(
        db: Session,
        source_id: int,
        target_id: int,
        relationship_type: str,
        strength: float = 1.0
    ) -> KnowledgeConnection:
        """Create a connection between two knowledge items."""
        # Verify both items exist
        source = db.query(KnowledgeItem).filter(KnowledgeItem.id == source_id).first()
        target = db.query(KnowledgeItem).filter(KnowledgeItem.id == target_id).first()
        
        if not source or not target:
            raise HTTPException(status_code=404, detail="One or both knowledge items not found")
        
        # Create connection
        connection = KnowledgeConnection(
            source_id=source_id,
            target_id=target_id,
            relationship_type=relationship_type,
            strength=strength
        )
        
        db.add(connection)
        db.commit()
        db.refresh(connection)
        return connection
    
    @staticmethod
    async def get_knowledge_graph(
        db: Session,
        user_id: int,
        repository_id: Optional[int] = None,
        central_item_id: Optional[int] = None,
        depth: int = 2
    ) -> Dict[str, Any]:
        """Get knowledge graph data for visualization."""
        # Get accessible repositories
        accessible_repos = await KnowledgeService.get_repositories(db, user_id)
        accessible_repo_ids = [repo.id for repo in accessible_repos]
        
        if repository_id and repository_id not in accessible_repo_ids:
            raise HTTPException(status_code=403, detail="You don't have access to this repository")
        
        # Filter by repository if specified
        repo_filter = KnowledgeItem.repository_id.in_(accessible_repo_ids)
        if repository_id:
            repo_filter = KnowledgeItem.repository_id == repository_id
        
        # Get all items and connections
        if central_item_id:
            # Start with the central item
            central_item = db.query(KnowledgeItem).filter(
                KnowledgeItem.id == central_item_id,
                repo_filter
            ).first()
            
            if not central_item:
                raise HTTPException(status_code=404, detail="Central knowledge item not found or not accessible")
            
            # Get connected items up to specified depth
            items, connections = KnowledgeService._get_connected_items(db, central_item, depth)
        else:
            # Get all items in the repository/repositories
            items = db.query(KnowledgeItem).filter(repo_filter).all()
            
            # Get all connections between these items
            item_ids = [item.id for item in items]
            connections = db.query(KnowledgeConnection).filter(
                KnowledgeConnection.source_id.in_(item_ids),
                KnowledgeConnection.target_id.in_(item_ids)
            ).all()
        
        # Format data for visualization
        nodes = []
        for item in items:
            nodes.append({
                "id": item.id,
                "label": item.title,
                "type": item.content_type,
                "importance": item.importance,
                "tags": [tag.name for tag in item.tags]
            })
        
        edges = []
        for conn in connections:
            edges.append({
                "source": conn.source_id,
                "target": conn.target_id,
                "type": conn.relationship_type,
                "strength": conn.strength
            })
        
        return {
            "nodes": nodes,
            "edges": edges
        }
    
    @staticmethod
    def _get_connected_items(db: Session, central_item: KnowledgeItem, depth: int) -> tuple:
        """Get items connected to the central item up to specified depth."""
        items = set([central_item])
        connections = set()
        
        # Process items level by level
        current_level = {central_item.id}
        for _ in range(depth):
            next_level = set()
            
            # Get all connections from current level
            for item_id in current_level:
                # Outgoing connections
                outgoing = db.query(KnowledgeConnection).filter(
                    KnowledgeConnection.source_id == item_id
                ).all()
                
                for conn in outgoing:
                    connections.add(conn)
                    target = db.query(KnowledgeItem).filter(KnowledgeItem.id == conn.target_id).first()
                    if target:
                        items.add(target)
                        next_level.add(target.id)
                
                # Incoming connections
                incoming = db.query(KnowledgeConnection).filter(
                    KnowledgeConnection.target_id == item_id
                ).all()
                
                for conn in incoming:
                    connections.add(conn)
                    source = db.query(KnowledgeItem).filter(KnowledgeItem.id == conn.source_id).first()
                    if source:
                        items.add(source)
                        next_level.add(source.id)
            
            # Update current level for next iteration
            current_level = next_level
            if not current_level:
                break
        
        return list(items), list(connections)


class DocumentService:
    """Service for managing document uploads and processing."""
    
    ALLOWED_EXTENSIONS = {'pdf', 'docx', 'txt', 'md', 'csv', 'json'}
    UPLOAD_DIR = "uploads/documents"
    
    @staticmethod
    async def upload_document(
        db: Session,
        file: UploadFile,
        uploader_id: int
    ) -> Document:
        """Upload a document and queue it for processing."""
        # Check file extension
        extension = file.filename.split('.')[-1].lower()
        if extension not in DocumentService.ALLOWED_EXTENSIONS:
            raise HTTPException(status_code=400, detail=f"File type not allowed. Allowed types: {', '.join(DocumentService.ALLOWED_EXTENSIONS)}")
        
        # Ensure upload directory exists
        os.makedirs(DocumentService.UPLOAD_DIR, exist_ok=True)
        
        # Generate unique filename
        timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
        unique_filename = f"{timestamp}_{file.filename}"
        file_path = os.path.join(DocumentService.UPLOAD_DIR, unique_filename)
        
        # Save file
        with open(file_path, "wb") as buffer:
            content = await file.read()
            buffer.write(content)
            file_size = len(content)
        
        # Create document record
        document = Document(
            filename=file.filename,
            file_path=file_path,
            file_type=extension,
            file_size=file_size,
            uploader_id=uploader_id
        )
        
        db.add(document)
        db.commit()
        db.refresh(document)
        
        # Queue document for processing
        # In a production system, this would be handled by a background task
        # For now, we'll process it synchronously
        await DocumentService.process_document(db, document.id)
        
        return document
    
    @staticmethod
    async def process_document(db: Session, document_id: int) -> Document:
        """Process a document and extract knowledge items."""
        document = db.query(Document).filter(Document.id == document_id).first()
        
        if not document:
            raise HTTPException(status_code=404, detail="Document not found")
        
        # Update status
        document.processing_status = "processing"
        db.commit()
        
        try:
            # Extract text based on file type
            text = await DocumentService._extract_text(document.file_path, document.file_type)
            
            # Create a default repository for the user if none exists
            from app.models.models import User
            user = db.query(User).filter(User.id == document.uploader_id).first()
            
            default_repo = db.query(KnowledgeRepository).filter(
                KnowledgeRepository.owner_id == document.uploader_id,
                KnowledgeRepository.name == f"{user.username}'s Knowledge"
            ).first()
            
            if not default_repo:
                default_repo = await KnowledgeService.create_repository(
                    db=db,
                    name=f"{user.username}'s Knowledge",
                    description="Default knowledge repository",
                    owner_id=document.uploader_id,
                    is_public=False
                )
            
            # Create knowledge item for the document
            knowledge_item = await KnowledgeService.create_knowledge_item(
                db=db,
                repository_id=default_repo.id,
                title=document.filename,
                content=text,
                creator_id=document.uploader_id,
                content_type="document",
                source_type="document",
                metadata={
                    "document_id": document.id,
                    "file_type": document.file_type,
                    "file_size": document.file_size
                }
            )
            
            # Link document to knowledge item
            document.knowledge_items.append(knowledge_item)
            
            # Update document status
            document.processed = True
            document.processing_status = "completed"
            db.commit()
            db.refresh(document)
            
            return document
        except Exception as e:
            logger.error(f"Error processing document: {e}")
            document.processing_status = "failed"
            db.commit()
            raise HTTPException(status_code=500, detail=f"Error processing document: {str(e)}")
    
    @staticmethod
    async def _extract_text(file_path: str, file_type: str) -> str:
        """Extract text from a document based on its type."""
        if file_type == 'txt' or file_type == 'md':
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        
        elif file_type == 'pdf':
            try:
                import PyPDF2
                text = ""
                with open(file_path, 'rb') as f:
                    pdf_reader = PyPDF2.PdfReader(f)
                    for page_num in range(len(pdf_reader.pages)):
                        text += pdf_reader.pages[page_num].extract_text() + "\n\n"
                return text
            except ImportError:
                return "PDF extraction requires PyPDF2 library. Please install it."
        
        elif file_type == 'docx':
            try:
                import docx
                doc = docx.Document(file_path)
                return "\n\n".join([para.text for para in doc.paragraphs])
            except ImportError:
                return "DOCX extraction requires python-docx library. Please install it."
        
        elif file_type == 'csv':
            try:
                import csv
                text = ""
                with open(file_path, 'r', encoding='utf-8') as f:
                    csv_reader = csv.reader(f)
                    for row in csv_reader:
                        text += ",".join(row) + "\n"
                return text
            except Exception as e:
                return f"Error extracting CSV: {str(e)}"
        
        elif file_type == 'json':
            try:
                import json
                with open(file_path, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                return json.dumps(data, indent=2)
            except Exception as e:
                return f"Error extracting JSON: {str(e)}"
        
        return f"Unsupported file type: {file_type}"


class WebResearchService:
    """Service for web research and integration with knowledge base."""
    
    @staticmethod
    async def search_web(
        db: Session,
        query: str,
        user_id: int,
        search_engine: str = "google",
        num_results: int = 5
    ) -> List[WebSearchResult]:
        """Perform a web search and store results."""
        try:
            # Use the info_search_web tool to get search results
            # In a real implementation, this would use an actual search API
            search_results = await WebResearchService._perform_search(query, search_engine, num_results)
            
            # Store search results
            db_results = []
            for result in search_results:
                web_result = WebSearchResult(
                    query=query,
                    url=result.get("url", ""),
                    title=result.get("title", ""),
                    snippet=result.get("snippet", ""),
                    search_engine=search_engine,
                    searcher_id=user_id
                )
                db.add(web_result)
                db_results.append(web_result)
            
            db.commit()
            for result in db_results:
                db.refresh(result)
            
            return db_results
        except Exception as e:
            logger.error(f"Error searching web: {e}")
            raise HTTPException(status_code=500, detail=f"Error searching web: {str(e)}")
    
    @staticmethod
    async def _perform_search(query: str, search_engine: str, num_results: int) -> List[Dict[str, str]]:
        """Perform a web search using the specified search engine."""
        # This is a mock implementation
        # In a real application, this would use an actual search API
        
        # Simulate search results
        mock_results = [
            {
                "url": f"https://example.com/result{i}",
                "title": f"Search Result {i} for '{query}'",
                "snippet": f"This is a snippet for search result {i} related to {query}..."
            }
            for i in range(1, num_results + 1)
        ]
        
        return mock_results
    
    @staticmethod
    async def add_to_knowledge_base(
        db: Session,
        web_result_id: int,
        repository_id: int,
        user_id: int
    ) -> KnowledgeItem:
        """Add a web search result to the knowledge base."""
        # Get web result
        web_result = db.query(WebSearchResult).filter(WebSearchResult.id == web_result_id).first()
        
        if not web_result:
            raise HTTPException(status_code=404, detail="Web search result not found")
        
        # Check repository access
        repository = await KnowledgeService.get_repository(db, repository_id, user_id)
        if not repository:
            raise HTTPException(status_code=404, detail="Repository not found or you don't have permission")
        
        # Fetch content if not already fetched
        if not web_result.content:
            web_result.content = await WebResearchService._fetch_web_content(web_result.url)
            db.commit()
        
        # Create knowledge item
        knowledge_item = await KnowledgeService.create_knowledge_item(
            db=db,
            repository_id=repository_id,
            title=web_result.title or f"Web content from {web_result.url}",
            content=web_result.content or web_result.snippet,
            creator_id=user_id,
            content_type="webpage",
            source_url=web_result.url,
            source_type="web",
            metadata={
                "search_query": web_result.query,
                "search_engine": web_result.search_engine,
                "search_date": web_result.search_date.isoformat() if web_result.search_date else None
            }
        )
        
        # Link web result to knowledge item
        web_result.added_to_knowledge = True
        web_result.knowledge_item_id = knowledge_item.id
        db.commit()
        
        return knowledge_item
    
    @staticmethod
    async def _fetch_web_content(url: str) -> str:
        """Fetch and extract content from a web page."""
        # This is a mock implementation
        # In a real application, this would use requests or similar to fetch the page
        
        # Simulate fetched content
        return f"This is the content fetched from {url}. In a real implementation, this would be the actual content of the web page, extracted and cleaned for use in the knowledge base."


class CitationService:
    """Service for managing citations of knowledge items by agents."""
    
    @staticmethod
    async def create_citation(
        db: Session,
        knowledge_item_id: int,
        agent_id: int,
        message_id: int,
        context: str,
        relevance_score: float = 0.8
    ) -> Citation:
        """Create a citation of a knowledge item by an agent."""
        # Verify knowledge item exists
        knowledge_item = db.query(KnowledgeItem).filter(KnowledgeItem.id == knowledge_item_id).first()
        if not knowledge_item:
            raise HTTPException(status_code=404, detail="Knowledge item not found")
        
        # Create citation
        citation = Citation(
            knowledge_item_id=knowledge_item_id,
            agent_id=agent_id,
            message_id=message_id,
            context=context,
            relevance_score=relevance_score
        )
        
        db.add(citation)
        db.commit()
        db.refresh(citation)
        return citation
    
    @staticmethod
    async def get_citations_by_agent(db: Session, agent_id: int) -> List[Citation]:
        """Get all citations by a specific agent."""
        return db.query(Citation).filter(Citation.agent_id == agent_id).all()
    
    @staticmethod
    async def get_citations_by_message(db: Session, message_id: int) -> List[Citation]:
        """Get all citations in a specific message."""
        return db.query(Citation).filter(Citation.message_id == message_id).all()
    
    @staticmethod
    async def get_citations_by_knowledge_item(db: Session, knowledge_item_id: int) -> List[Citation]:
        """Get all citations of a specific knowledge item."""
        return db.query(Citation).filter(Citation.knowledge_item_id == knowledge_item_id).all()
